from docx import Document
import os

def create_resume(filename):
    doc = Document()
    
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer | john.doe@email.com | +1 123 456 7890')
    doc.add_paragraph('LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe')
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Full Stack Developer with 5 years of experience in building scalable web applications. Expert in Python, JavaScript, and Cloud Computing. Proven track record of leading teams and delivering high-quality software.')
    
    doc.add_heading('Experience', level=1)
    doc.add_heading('Senior Software Engineer - Tech Solutions Inc. (2020 - Present)', level=2)
    doc.add_paragraph('- Developed and maintained full stack applications using React and Node.js.')
    doc.add_paragraph('- Improved system performance by 40% through code optimization.')
    doc.add_paragraph('- Managed a team of 5 developers and led multiple successful product launches.')
    
    doc.add_heading('Software Engineer - Web Innovations (2018 - 2020)', level=2)
    doc.add_paragraph('- Built responsive frontends using HTML, CSS, and JavaScript.')
    doc.add_paragraph('- Integrated backend services with REST APIs.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, JavaScript, React, Node.js, SQL, MongoDB, AWS, Docker, Git, HTML, CSS, REST APIs, Java, Docker, Kubernetes')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Technology in Computer Science - University of Technology (2018)')
    
    doc.save(filename)

if __name__ == "__main__":
    create_resume("test_resume.docx")
    print("test_resume.docx created successfully.")
