"""
DOCX Resume Reader - Extracts text from Word documents using python-docx.
"""
from docx import Document
import io


def extract_text_from_docx(file_path=None, file_stream=None):
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path: Path to DOCX file on disk
        file_stream: File-like object (for uploaded files)
    
    Returns:
        str: Extracted text content
    """
    text = ""
    try:
        if file_stream:
            doc = Document(file_stream)
        elif file_path:
            doc = Document(file_path)
        else:
            return ""

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"

    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

    return text.strip()
